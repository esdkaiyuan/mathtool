from __future__ import annotations

import re
import sys
from copy import deepcopy
from pathlib import Path, PurePosixPath

import latex2mathml.converter
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from lxml import etree


MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M = f"{{{MATH_NS}}}"
W = f"{{{WORD_NS}}}"
RELATION_BOUNDARY = re.compile(r"[=<>≤≥≈≠]")
SUPPORTED_FORMULA_FONTS = (
    "Cambria Math",
    "Times New Roman",
    "Latin Modern Math",
    "STIX Two Text",
    "Arial",
)

OFFICE_MML2OMML_PATHS = (
    Path(r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL"),
    Path(r"C:\Program Files (x86)\Microsoft Office\root\Office16\MML2OMML.XSL"),
    PurePosixPath("/Applications/Microsoft Word.app/Contents/Resources/MML2OMML.XSL"),
    PurePosixPath("/Applications/Microsoft Word.app/Contents/Resources/Microsoft/Office/MML2OMML.XSL"),
)


def find_mml2omml_path() -> Path:
    bundled = Path(getattr(sys, "_MEIPASS", Path(__file__).resolve().parent)) / "MML2OMML.XSL"
    if bundled.exists():
        return bundled

    import os

    configured = os.environ.get("MML2OMML_XSL")
    if configured and Path(configured).exists():
        return Path(configured)

    for path in OFFICE_MML2OMML_PATHS:
        if path.exists():
            return path

    raise FileNotFoundError(
        "未找到 Microsoft Office 的 MML2OMML.XSL。请确认本机已安装 Microsoft Word。"
    )


def normalize_latex(latex: str) -> str:
    value = (latex or "").strip()
    value = re.sub(r"^\\\(|\\\)$", "", value)
    value = re.sub(r"^\\\[|\\\]$", "", value)
    value = value.strip("$")
    return value.strip()


def _is_empty_math_element(element: etree._Element | None) -> bool:
    if element is None:
        return True
    return len(element) == 0 and not (element.text or "").strip()


def _set_math_run_text(run: etree._Element, text: str) -> etree._Element:
    text_node = run.find(f"{M}t")
    if text_node is None:
        text_node = etree.SubElement(run, f"{M}t")
    text_node.text = text
    for extra in run.findall(f"{M}t")[1:]:
        extra.getparent().remove(extra)
    return run


def _split_at_relation_boundary(element: etree._Element) -> tuple[etree._Element | None, etree._Element | None]:
    if element.tag != f"{M}r":
        return element, None

    text_nodes = element.findall(f"{M}t")
    if len(text_nodes) != 1 or text_nodes[0].text is None:
        return element, None

    text = text_nodes[0].text
    match = RELATION_BOUNDARY.search(text)
    if match is None:
        return element, None

    if match.start() == 0:
        return None, element

    before = _set_math_run_text(deepcopy(element), text[: match.start()])
    after = _set_math_run_text(deepcopy(element), text[match.start() :])
    return before, after


def repair_empty_nary_operands(omml_root: etree._Element) -> etree._Element:
    for nary in list(omml_root.iter(f"{M}nary")):
        operand = nary.find(f"{M}e")
        if not _is_empty_math_element(operand):
            continue

        parent = nary.getparent()
        if parent is None or operand is None:
            continue

        moved_any = False
        while True:
            siblings = list(parent)
            try:
                nary_index = siblings.index(nary)
            except ValueError:
                break

            next_index = nary_index + 1
            if next_index >= len(siblings):
                break

            sibling = siblings[next_index]
            if sibling.tag == f"{M}nary":
                break

            before, after = _split_at_relation_boundary(sibling)
            if before is None:
                break

            parent.remove(sibling)
            operand.append(before)
            moved_any = True

            if after is not None:
                parent.insert(next_index, after)
                break

        if not moved_any:
            continue

    return omml_root


def _set_word_run_font(run_properties: etree._Element, font_name: str) -> None:
    fonts = run_properties.find(f"{W}rFonts")
    if fonts is None:
        fonts = etree.Element(f"{W}rFonts")
        run_properties.insert(0, fonts)

    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(f"{W}{attribute}", font_name)


def _ensure_word_run_properties(math_run: etree._Element) -> etree._Element:
    run_properties = math_run.find(f"{W}rPr")
    if run_properties is not None:
        return run_properties

    run_properties = etree.Element(f"{W}rPr")
    math_run_properties = math_run.find(f"{M}rPr")
    insert_index = 1 if math_run_properties is not None else 0
    math_run.insert(insert_index, run_properties)
    return run_properties


def _ensure_control_properties(math_properties: etree._Element) -> etree._Element:
    control_properties = math_properties.find(f"{M}ctrlPr")
    if control_properties is None:
        control_properties = etree.SubElement(math_properties, f"{M}ctrlPr")

    run_properties = control_properties.find(f"{W}rPr")
    if run_properties is None:
        run_properties = etree.SubElement(control_properties, f"{W}rPr")

    return run_properties


def apply_formula_font(omml_root: etree._Element, font_name: str | None) -> etree._Element:
    selected_font = (font_name or "").strip()
    if not selected_font:
        return omml_root

    for math_run in omml_root.iter(f"{M}r"):
        _set_word_run_font(_ensure_word_run_properties(math_run), selected_font)

    property_tags = {
        f"{M}barPr",
        f"{M}borderBoxPr",
        f"{M}boxPr",
        f"{M}dPr",
        f"{M}eqArrPr",
        f"{M}fPr",
        f"{M}funcPr",
        f"{M}groupChrPr",
        f"{M}limLowPr",
        f"{M}limUppPr",
        f"{M}mPr",
        f"{M}naryPr",
        f"{M}phantPr",
        f"{M}radPr",
        f"{M}sPrePr",
        f"{M}sSubPr",
        f"{M}sSubSupPr",
        f"{M}sSupPr",
    }
    for math_properties in omml_root.iter():
        if math_properties.tag in property_tags:
            _set_word_run_font(_ensure_control_properties(math_properties), selected_font)

    return omml_root


def latex_to_omml(latex: str, formula_font: str | None = None) -> etree._Element:
    cleaned = normalize_latex(latex)
    if not cleaned:
        raise ValueError("公式不能为空")

    mathml = latex2mathml.converter.convert(cleaned)
    mathml_root = etree.fromstring(mathml.encode("utf-8"))
    transform = etree.XSLT(etree.parse(str(find_mml2omml_path())))
    omml_tree = transform(mathml_root)
    omml_root = omml_tree.getroot()

    if omml_root is None:
        raise ValueError("公式转换失败，未生成 Word OMML。")

    repair_empty_nary_operands(omml_root)
    apply_formula_font(omml_root, formula_font)
    return omml_root


def _set_cell_text_style(run, size_pt: int = 11) -> None:
    run.font.name = "Microsoft YaHei UI"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei UI")
    run.font.size = Pt(size_pt)


def add_omml_to_paragraph(paragraph, omml: etree._Element) -> None:
    paragraph._p.append(deepcopy(omml))


def build_docx(
    output_path: str | Path,
    title: str,
    intro: str,
    equations: list[str],
    formula_font: str | None = None,
) -> Path:
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei UI"
    styles["Normal"].font.size = Pt(11)
    styles["Title"].font.name = "Microsoft YaHei UI"

    section = doc.sections[0]
    section.top_margin = Pt(56)
    section.bottom_margin = Pt(56)
    section.left_margin = Pt(64)
    section.right_margin = Pt(64)

    heading = doc.add_paragraph()
    heading.style = "Title"
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_run = heading.add_run(title.strip() or "Word 原生公式")
    heading_run.font.name = "Microsoft YaHei UI"
    heading_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei UI")

    intro_text = intro.strip()
    if intro_text:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(intro_text)
        _set_cell_text_style(run)

    valid_equations = [normalize_latex(item) for item in equations if normalize_latex(item)]
    if not valid_equations:
        valid_equations = [r"E = mc^2"]

    for index, equation in enumerate(valid_equations, start=1):
        label = doc.add_paragraph()
        label_run = label.add_run(f"公式 {index}")
        label_run.bold = True
        _set_cell_text_style(label_run, 10)

        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(14)
        add_omml_to_paragraph(paragraph, latex_to_omml(equation, formula_font=formula_font))

    note = doc.add_paragraph()
    note_run = note.add_run("提示：以上公式为 Word 原生公式对象，可在 Word 中双击编辑，也可复制到其他 Word 文档。")
    _set_cell_text_style(note_run, 9)

    doc.save(output)
    return output


def parse_equations(text: str) -> list[str]:
    lines = []
    for raw_line in (text or "").splitlines():
        line = raw_line.strip()
        if line:
            lines.append(line)
    return lines
